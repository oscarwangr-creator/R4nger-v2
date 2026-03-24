"""
RTF v2.0 — Reports: DOCX Reporter
Generates a professional Word document report.
Requires: pip install python-docx
"""
from __future__ import annotations
import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

SEV_COLORS_HEX = {
    "critical": "DC2626",
    "high":     "EA580C",
    "medium":   "D97706",
    "low":      "65A30D",
    "info":     "0284C7",
}

class DocxReporter:
    def generate(self, data: Dict[str, Any], output_path: str) -> str:
        p = Path(output_path); p.parent.mkdir(parents=True, exist_ok=True)
        try:
            return self._generate_docx(data, p)
        except ImportError:
            # Plain text fallback if python-docx not installed
            self._write_text_fallback(data, p)
            return str(p.resolve())

    def _generate_docx(self, data: Dict[str, Any], p: Path) -> str:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)

        findings = data.get("findings", [])
        target   = str(data.get("target",""))
        profile  = str(data.get("profile",""))
        pid      = str(data.get("pipeline_id",""))
        operator = str(data.get("operator","RTF Operator"))
        gen_at   = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
        entities = data.get("entities",{})
        ai_anal  = data.get("ai_analysis",{})

        def add_heading(text, level=1, color_hex="1E293B"):
            h = doc.add_heading(text, level=level)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in h.runs:
                run.font.color.rgb = RGBColor.from_string(color_hex)
            return h

        def add_color_run(para, text, color_hex, bold=False):
            run = para.add_run(text)
            run.font.color.rgb = RGBColor.from_string(color_hex)
            run.bold = bold
            return run

        def sev_color(sev: str) -> str:
            return SEV_COLORS_HEX.get(sev.lower(), "64748B")

        # ── Cover ────────────────────────────────────────────────
        doc.add_paragraph()
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("RTF RedTeam Report")
        title_run.font.size  = Pt(28)
        title_run.bold       = True
        title_run.font.color.rgb = RGBColor.from_string("1E3A5F")

        doc.add_paragraph()
        meta_lines = [
            ("Target:",    target),
            ("Profile:",   profile),
            ("Pipeline:",  pid),
            ("Operator:",  operator),
            ("Generated:", gen_at),
        ]
        for label, val in meta_lines:
            p_meta = doc.add_paragraph()
            p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_l = p_meta.add_run(label + "  ")
            run_l.bold = True; run_l.font.color.rgb = RGBColor.from_string("64748B")
            run_v = p_meta.add_run(val)
            run_v.font.size = Pt(11)

        doc.add_page_break()

        # ── Executive Summary ────────────────────────────────────
        add_heading("Executive Summary", level=1)

        from collections import Counter
        sev_counts = Counter(str(f.get("severity","info")).lower() for f in findings)
        risk_score = (sev_counts["critical"]*10 + sev_counts["high"]*5 +
                      sev_counts["medium"]*2 + sev_counts["low"])

        # Summary table
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr_cells = tbl.rows[0].cells
        for c, txt in zip(hdr_cells, ["Severity", "Count", "Score"]):
            c.text = txt
            for run in c.paragraphs[0].runs:
                run.bold = True

        for sev in ("critical","high","medium","low","info"):
            row_cells = tbl.add_row().cells
            row_cells[0].text = sev.upper()
            row_cells[1].text = str(sev_counts.get(sev, 0))
            w = {"critical":10,"high":5,"medium":2,"low":1,"info":0}
            row_cells[2].text = str(sev_counts.get(sev,0) * w.get(sev,0))
            for para in row_cells[0].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor.from_string(sev_color(sev))
                    run.bold = True

        total_row = tbl.add_row().cells
        total_row[0].text = "TOTAL"
        total_row[1].text = str(len(findings))
        total_row[2].text = str(risk_score)
        for cell in total_row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        doc.add_paragraph()
        p_risk = doc.add_paragraph()
        p_risk.add_run("Overall Risk Score: ").bold = True
        add_color_run(p_risk, str(risk_score),
                      "DC2626" if risk_score>=20 else "EA580C" if risk_score>=10 else "D97706" if risk_score>=5 else "65A30D",
                      bold=True)

        # ── AI Analysis ──────────────────────────────────────────
        if ai_anal:
            doc.add_paragraph()
            add_heading("AI Analysis", level=2)
            conf = ai_anal.get("confidence_score", "N/A")
            rl   = ai_anal.get("risk_level", "")
            summ = str(ai_anal.get("executive_summary", ai_anal.get("identity_summary","")))[:600]
            p_ai = doc.add_paragraph()
            p_ai.add_run(f"Confidence: {conf}/100   Risk Level: ").bold = True
            p_ai.add_run(rl)
            if summ:
                doc.add_paragraph(summ)

        # ── Findings ─────────────────────────────────────────────
        doc.add_page_break()
        add_heading(f"Findings ({len(findings)} total)", level=1)

        sorted_f = sorted(findings, key=lambda f: {"critical":0,"high":1,"medium":2,"low":3,"info":4}.get(
            str(f.get("severity","info")).lower(), 4))

        for f in sorted_f:
            sev   = str(f.get("severity","info")).lower()
            title = str(f.get("title",""))[:120]
            desc  = str(f.get("description",""))[:400]
            tgt   = str(f.get("target",""))
            tags  = " | ".join(str(t) for t in f.get("tags",[])[:5])

            p_find = doc.add_paragraph()
            add_color_run(p_find, f"[{sev.upper()}] ", sev_color(sev), bold=True)
            p_find.add_run(title).bold = True

            if tgt:
                doc.add_paragraph(f"Target: {tgt}").style = "No Spacing"
            if desc:
                doc.add_paragraph(desc).style = "No Spacing"
            if tags:
                p_tags = doc.add_paragraph()
                p_tags.add_run("Tags: ").bold = True
                p_tags.add_run(tags)
            doc.add_paragraph()

        # ── Entities ─────────────────────────────────────────────
        if entities:
            doc.add_page_break()
            add_heading("Discovered Entities", level=1)
            for etype, values in entities.items():
                if not values: continue
                add_heading(etype.title(), level=2)
                vlist = list(values)[:50] if hasattr(values,"__iter__") and not isinstance(values,str) else [values]
                for v in vlist:
                    doc.add_paragraph(str(v), style="List Bullet")

        doc.save(str(p))
        return str(p.resolve())

    def _write_text_fallback(self, data: Dict[str, Any], p: Path) -> None:
        """Plain text fallback when python-docx is not installed."""
        lines = [
            "RTF RedTeam Report",
            "=" * 60,
            f"Target:    {data.get('target','')}",
            f"Profile:   {data.get('profile','')}",
            f"Generated: {datetime.utcnow().isoformat()}",
            "",
            "NOTE: Install python-docx for proper DOCX output:",
            "  pip install python-docx",
            "",
            "FINDINGS:",
            "-" * 60,
        ]
        for f in data.get("findings", []):
            lines.append(f"[{str(f.get('severity','info')).upper()}] {f.get('title','')}")
            if f.get("description"):
                lines.append(f"  {f.get('description','')[:200]}")
            lines.append("")
        p.with_suffix(".txt").write_text("\n".join(lines), encoding="utf-8")
        p.write_text("\n".join(lines), encoding="utf-8")
